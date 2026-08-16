from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from typing import List, Optional
from app.models.exam import Exam, ExamQuestion, ExamChoice, ExamAttempt, ExamAnswer
from app.schemas.exam import ExamCreate, ExamResponse, ExamAttemptCreate, ExamAttemptResponse, ExamAttemptSubmit, ExamAttemptDetailResponse, ExamReviewResponse
from app.utils.auth import get_current_user
from app.models.user import User
from app.models.student import Student
from app.models.course import Course
from app.models.submission import Submission
import PyPDF2
import docx
import io
import re
import random
from datetime import datetime

exam_router = APIRouter(prefix="/exams", tags=["Exams"])

def format_exam_response(exam: Exam, course: Optional[dict] = None) -> dict:
    data = exam.model_dump()
    data["id"] = exam.int_id
    if course:
        data["course"] = course
    for q in data.get("questions", []):
        q["id"] = q.get("int_id", 0)
        for c in q.get("choices", []):
            c["id"] = c.get("int_id", 0)
    return data



@exam_router.get("/stats")
async def get_exam_stats():
    total_exams = await Exam.find_all().count()
    
    # Submissions today
    today = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    
    from app.models.student import Student
    
    exams = await Exam.find_all().to_list()
    
    submissions_today = 0
    total_attempts = 0
    total_expected_attempts = 0
    
    for exam in exams:
        student_count = await Student.find(Student.course_id == exam.course_id).count()
        total_expected_attempts += student_count if student_count > 0 else 1
        
        for attempt in exam.attempts:
            if attempt.status == "submitted":
                total_attempts += 1
                if attempt.end_time and attempt.end_time >= today:
                    submissions_today += 1
    
    if total_expected_attempts > 0:
        avg = min(100.0, (total_attempts / total_expected_attempts) * 100)
    else:
        avg = 0.0
        
    avg_completion = f"{avg:.1f}%"
    
    return {
        "total_exams": total_exams,
        "submissions_today": submissions_today,
        "avg_completion": avg_completion,
        "pending_ai_review": sum(1 for submission in await Submission.find(Submission.grade == None).to_list())
    }

@exam_router.get("/", response_model=List[ExamResponse])
async def get_all_exams(course_id: Optional[int] = None, current_user: User = Depends(get_current_user)):
    if current_user.role == "student":
        # Find the student record to get the course_id
        student = await Student.find_one(Student.email == current_user.email)
        if not student:
            return []
        # Students only see published exams for their course
        exams = await Exam.find(
            Exam.course_id == student.course_id,
            Exam.status == "published"
        ).to_list()
        
        valid_exams = []
        for e in exams:
            completed_attempts = sum(1 for a in e.attempts if a.student_id == current_user.int_id and a.status != "in_progress")
            if completed_attempts < e.attempts_allowed:
                valid_exams.append(e)
                
        courses = await Course.find({"int_id": {"$in": [e.course_id for e in valid_exams]}}).to_list()
        course_map = {c.int_id: {"id": c.int_id, "name": c.name, "code": c.code} for c in courses}
        return [ExamResponse(**format_exam_response(e, course_map.get(e.course_id))) for e in valid_exams]
    
    # Teachers/Admins can see all or filter by course_id
    query = Exam.find_all()
    if course_id:
        query = query.find(Exam.course_id == course_id)
        
    exams = await query.to_list()
    courses = await Course.find({"int_id": {"$in": [e.course_id for e in exams]}}).to_list()
    course_map = {c.int_id: {"id": c.int_id, "name": c.name, "code": c.code} for c in courses}
    return [ExamResponse(**format_exam_response(e, course_map.get(e.course_id))) for e in exams]

@exam_router.post("/", response_model=ExamResponse)
async def create_exam(exam_data: ExamCreate):
    new_exam = Exam(
        course_id=exam_data.course_id,
        title=exam_data.title,
        description=exam_data.description,
        time_limit=exam_data.time_limit,
        attempts_allowed=exam_data.attempts_allowed,
        randomize_questions=exam_data.randomize_questions,
        status=exam_data.status
    )
    await new_exam.assign_id()

    question_id_counter = 1
    choice_id_counter = 1
    
    for i, q_data in enumerate(exam_data.questions):
        new_question = ExamQuestion(
            int_id=question_id_counter,
            question_text=q_data.question_text,
            question_type=q_data.question_type,
            points=q_data.points,
            order=q_data.order or i
        )
        question_id_counter += 1

        for choice_data in q_data.choices:
            new_choice = ExamChoice(
                int_id=choice_id_counter,
                choice_text=choice_data.choice_text,
                is_correct=choice_data.is_correct
            )
            choice_id_counter += 1
            new_question.choices.append(new_choice)
            
        new_exam.questions.append(new_question)
    
    await new_exam.insert()
    return ExamResponse(**format_exam_response(new_exam))

@exam_router.get("/course/{course_id}", response_model=List[ExamResponse])
async def get_course_exams(course_id: int):
    exams = await Exam.find(Exam.course_id == course_id).to_list()
    return [ExamResponse(**format_exam_response(e)) for e in exams]

@exam_router.get("/template")
async def download_exam_template():
    """Download a Kahoot-style Excel template for bulk question import."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from fastapi.responses import StreamingResponse
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Questions"
    
    # ── Styling ──
    header_font = Font(name="Calibri", bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="264796", end_color="264796", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="D0D0D0"),
        right=Side(style="thin", color="D0D0D0"),
        top=Side(style="thin", color="D0D0D0"),
        bottom=Side(style="thin", color="D0D0D0"),
    )
    
    # ── Headers ──
    headers = ["Question", "Option A", "Option B", "Option C", "Option D", "Correct Answer (A/B/C/D)", "Points"]
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border
    
    # ── Column widths ──
    ws.column_dimensions["A"].width = 50  # Question
    ws.column_dimensions["B"].width = 25  # Option A
    ws.column_dimensions["C"].width = 25  # Option B
    ws.column_dimensions["D"].width = 25  # Option C
    ws.column_dimensions["E"].width = 25  # Option D
    ws.column_dimensions["F"].width = 18  # Correct Answer
    ws.column_dimensions["G"].width = 10  # Points
    
    # ── Example rows ──
    example_data = [
        ["What is the capital of France?", "London", "Paris", "Berlin", "Madrid", "B", 1],
        ["Which planet is known as the Red Planet?", "Venus", "Jupiter", "Mars", "Saturn", "C", 1],
    ]
    
    example_fill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
    example_font = Font(name="Calibri", size=11, color="666666", italic=True)
    
    for row_idx, row_data in enumerate(example_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.fill = example_fill
            cell.font = example_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    
    # ── Instructions sheet ──
    ws_info = wb.create_sheet("Instructions")
    instructions = [
        "HOW TO USE THIS TEMPLATE",
        "",
        "1. Fill in your questions starting from Row 2 on the 'Questions' sheet.",
        "2. You can delete or overwrite the example rows.",
        "3. Each row = one MCQ question.",
        "4. Fill in columns A through E with the question and four options.",
        "5. In column F ('Correct Answer'), type A, B, C, or D to indicate the correct option.",
        "6. Column G ('Points') is optional — defaults to 1 if left blank.",
        "7. Save the file and upload it back in the Exam Creator.",
        "",
        "RULES:",
        "• All questions must have at least Options A and B filled.",
        "• The Correct Answer column must contain A, B, C, or D.",
        "• Do not modify the header row.",
    ]
    
    for row_idx, line in enumerate(instructions, 1):
        cell = ws_info.cell(row=row_idx, column=1, value=line)
        if row_idx == 1:
            cell.font = Font(name="Calibri", bold=True, size=14, color="264796")
        elif line.startswith("RULES"):
            cell.font = Font(name="Calibri", bold=True, size=11)
        else:
            cell.font = Font(name="Calibri", size=11)
    ws_info.column_dimensions["A"].width = 80
    
    # ── Write to buffer ──
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=exam_questions_template.xlsx"}
    )


@exam_router.post("/import-excel")
async def import_exam_from_excel(file: UploadFile = File(...)):
    """Import MCQ questions from a filled Excel template."""
    import openpyxl
    
    if not file.filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Please upload an Excel file (.xlsx)")
    
    contents = await file.read()
    
    try:
        wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read Excel file: {str(e)}")
    
    ws = wb.active
    
    # Validate headers
    headers = [str(cell.value or "").strip().lower() for cell in ws[1]]
    if len(headers) < 6:
        raise HTTPException(
            status_code=400,
            detail="Invalid template format. Please use the provided template with at least 6 columns."
        )
    
    questions = []
    errors = []
    
    for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        # Skip completely empty rows
        if not row or all(cell is None or str(cell).strip() == "" for cell in row):
            continue
        
        question_text = str(row[0] or "").strip() if len(row) > 0 else ""
        option_a = str(row[1] or "").strip() if len(row) > 1 else ""
        option_b = str(row[2] or "").strip() if len(row) > 2 else ""
        option_c = str(row[3] or "").strip() if len(row) > 3 else ""
        option_d = str(row[4] or "").strip() if len(row) > 4 else ""
        correct_raw = str(row[5] or "").strip().upper() if len(row) > 5 else ""
        points_raw = row[6] if len(row) > 6 else None
        
        # Validate question
        if not question_text:
            errors.append(f"Row {row_idx}: Missing question text")
            continue
        
        if not option_a or not option_b:
            errors.append(f"Row {row_idx}: At least Options A and B are required")
            continue
        
        if correct_raw not in ("A", "B", "C", "D"):
            errors.append(f"Row {row_idx}: Correct Answer must be A, B, C, or D (got '{correct_raw}')")
            continue
        
        # Build choices
        options = [
            {"label": "A", "text": option_a},
            {"label": "B", "text": option_b},
        ]
        if option_c:
            options.append({"label": "C", "text": option_c})
        if option_d:
            options.append({"label": "D", "text": option_d})
        
        # Validate correct answer points to an existing option
        if correct_raw == "C" and not option_c:
            errors.append(f"Row {row_idx}: Correct answer is C but Option C is empty")
            continue
        if correct_raw == "D" and not option_d:
            errors.append(f"Row {row_idx}: Correct answer is D but Option D is empty")
            continue
        
        try:
            points = float(points_raw) if points_raw is not None else 1.0
        except (ValueError, TypeError):
            points = 1.0
        
        questions.append({
            "question_text": question_text,
            "points": points,
            "choices": [
                {
                    "choice_text": o["text"],
                    "is_correct": (o["label"] == correct_raw)
                }
                for o in options
            ]
        })
    
    if not questions and errors:
        raise HTTPException(status_code=400, detail=f"No valid questions found. Errors: {'; '.join(errors[:5])}")
    
    return {
        "questions": questions,
        "imported_count": len(questions),
        "errors": errors[:10] if errors else []
    }

@exam_router.get("/{exam_id}", response_model=ExamResponse)
async def get_exam(exam_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    return ExamResponse(**format_exam_response(exam))

@exam_router.post("/{exam_id}/start", response_model=ExamAttemptResponse)
async def start_exam_attempt(exam_id: int, current_user: User = Depends(get_current_user)):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Check for existing in-progress attempt to allow resuming
    existing_attempt = next((a for a in exam.attempts if a.student_id == current_user.int_id and a.status == "in_progress"), None)
    
    if existing_attempt:
        res = existing_attempt.model_dump()
        res["id"] = existing_attempt.int_id
        res["exam_id"] = exam.int_id
        return res
    
    # Check completed attempts
    completed_attempts = sum(1 for a in exam.attempts if a.student_id == current_user.int_id and a.status != "in_progress")
    
    if completed_attempts >= exam.attempts_allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Maximum attempts reached")

    attempt_id = max([a.int_id for a in exam.attempts] + [0]) + 1
    
    new_attempt = ExamAttempt(
        int_id=attempt_id,
        student_id=current_user.int_id,
        status="in_progress"
    )
    
    exam.attempts.append(new_attempt)
    await exam.save()
    
    res = new_attempt.model_dump()
    res["id"] = new_attempt.int_id
    res["exam_id"] = exam.int_id
    return res

@exam_router.post("/{exam_id}/attempts/{attempt_id}/submit", response_model=ExamAttemptResponse)
async def submit_exam_attempt(exam_id: int, attempt_id: int, submission: ExamAttemptSubmit):
    # Find the exam containing this attempt
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    attempt_idx = next((i for i, a in enumerate(exam.attempts) if a.int_id == attempt_id), -1)
    if attempt_idx == -1:
        raise HTTPException(status_code=404, detail="Attempt not found")
        
    attempt = exam.attempts[attempt_idx]
    
    if attempt.status != "in_progress":
        raise HTTPException(status_code=400, detail="Attempt already submitted or invalid")

    total_score = 0.0
    answer_id_counter = max([ans.int_id for ans in attempt.answers] + [0]) + 1
    
    for ans in submission.answers:
        new_answer = ExamAnswer(
            int_id=answer_id_counter,
            attempt_int_id=attempt_id,
            question_int_id=ans.question_id,
            selected_choice_id=ans.selected_choice_id
        )
        answer_id_counter += 1
        attempt.answers.append(new_answer)
        
        # Grading
        if ans.selected_choice_id:
            question = next((q for q in exam.questions if q.int_id == ans.question_id), None)
            if question:
                choice = next((c for c in question.choices if c.int_id == ans.selected_choice_id), None)
                if choice and choice.is_correct:
                    total_score += (question.points or 0.0)

    attempt.score = total_score
    attempt.status = "submitted"
    attempt.end_time = datetime.utcnow()
    
    await exam.save()
    
    res = attempt.model_dump()
    res["id"] = attempt.int_id
    res["exam_id"] = exam.int_id
    return res

@exam_router.post("/extract")
async def extract_exam_questions(file: UploadFile = File(...)):
    """Extract MCQ questions from PDF/DOCX using AI (Groq) with regex fallback."""
    contents = await file.read()
    text = ""
    
    if file.filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(contents))
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(contents))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "  ".join([cell.text.strip() for cell in row.cells])
                text_parts.append(row_text)
        text = "\n".join(text_parts)
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Please upload a .pdf or .docx file.")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text from the uploaded file.")

    # ── Attempt 1: AI-powered extraction via Groq ──
    try:
        from app.services.groq_service import GroqService, DEFAULT_MODEL
        if GroqService._available and GroqService._client:
            # Truncate very long documents to avoid token limits
            truncated_text = text[:12000] if len(text) > 12000 else text

            prompt = f"""You are an expert exam parser. Extract ALL multiple-choice questions (MCQs) from the following exam document text.

For each question, return:
- question_text: the full question text
- choices: an array of objects, each with "choice_text" (the option text) and "is_correct" (boolean, true for the correct answer)

Rules:
1. Extract EVERY question you can find — do not skip any.
2. Each question MUST have at least 2 choices (ideally 4).
3. If the correct answer is indicated anywhere in the document (answer key, inline marking, bold/underlined option), set is_correct=true for that choice. If NO correct answer is identifiable, set all is_correct=false.
4. Clean up the text — remove question numbers, trailing whitespace, and artifacts from PDF extraction.
5. Return ONLY a valid JSON array. No markdown fences, no explanation, no extra text.

Example output format:
[
  {{
    "question_text": "What is the capital of France?",
    "choices": [
      {{"choice_text": "London", "is_correct": false}},
      {{"choice_text": "Paris", "is_correct": true}},
      {{"choice_text": "Berlin", "is_correct": false}},
      {{"choice_text": "Madrid", "is_correct": false}}
    ]
  }}
]

Document text:
\"\"\"
{truncated_text}
\"\"\"

Return ONLY the JSON array:"""

            response = GroqService._client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=DEFAULT_MODEL,
                max_tokens=4096,
                temperature=0.1,
            )
            
            ai_text = response.choices[0].message.content.strip()
            
            # Strip markdown code fences if present
            if ai_text.startswith("```"):
                ai_text = re.sub(r'^```(?:json)?\s*', '', ai_text)
                ai_text = re.sub(r'\s*```$', '', ai_text)
            
            import json
            questions = json.loads(ai_text)
            
            if isinstance(questions, list) and len(questions) > 0:
                # Validate structure
                valid_questions = []
                for q in questions:
                    if isinstance(q, dict) and "question_text" in q and "choices" in q:
                        if isinstance(q["choices"], list) and len(q["choices"]) >= 2:
                            valid_choices = [
                                {
                                    "choice_text": str(c.get("choice_text", "")).strip(),
                                    "is_correct": bool(c.get("is_correct", False))
                                }
                                for c in q["choices"] if isinstance(c, dict)
                            ]
                            if len(valid_choices) >= 2:
                                valid_questions.append({
                                    "question_text": str(q["question_text"]).strip(),
                                    "choices": valid_choices
                                })
                
                if valid_questions:
                    return valid_questions
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"AI extraction failed, falling back to regex: {e}")

    # ── Attempt 2: Improved regex-based fallback ──
    # Try multiple question patterns for broader compatibility
    patterns_to_try = [
        # Pattern 1: Explicit "Q" or "Question" prefix
        r'(?:^|\n)\s*(?:Q(?:uestion)?)\s*(\d+)\s*[\.)\:\-]?\s*',
        # Pattern 2: Numbered questions (e.g. "1.", "1)", "1:")
        r'(?:^|\n)\s*(\d+)\s*[\.)\:\-]\s+',
    ]
    
    questions = []
    
    for q_pattern in patterns_to_try:
        q_splits = re.split(q_pattern, text, flags=re.IGNORECASE)
        if len(q_splits) > 2:  # Found matches (split produces alternating groups)
            # Process blocks — every other element is a question block
            for i in range(1, len(q_splits), 2):
                block = q_splits[i + 1] if (i + 1) < len(q_splits) else ""
                if not block.strip():
                    continue
                
                # Extract inline correct answer if present
                correct_answer = ""
                ans_match = re.search(
                    r'(?:^|\n|\s)(?:Answer|Ans|Correct)[^A-Za-z]*([A-Ea-e])\b',
                    block, re.IGNORECASE
                )
                if ans_match:
                    correct_answer = ans_match.group(1).upper()
                    block = block[:ans_match.start()] + block[ans_match.end():]
                
                # Find options using multiple option patterns
                option_patterns = [
                    # (A) text or A) text or A. text
                    r'(?:^|\n)\s*\(?([A-Ea-e])\)?[\.\)]\s*(.*?)(?=(?:\n\s*\(?[A-Ea-e]\)?[\.\)]|\Z))',
                    # a) text, a. text (lowercase)
                    r'(?:^|\n)\s*\(?([a-e])\)?[\.\)]\s*(.*?)(?=(?:\n\s*\(?[a-e]\)?[\.\)]|\Z))',
                ]
                
                options = []
                for opt_pattern in option_patterns:
                    opt_matches = re.findall(opt_pattern, block, re.DOTALL | re.IGNORECASE)
                    if opt_matches and len(opt_matches) >= 2:
                        for label, opt_text in opt_matches:
                            clean_text = opt_text.strip().replace('\n', ' ')
                            if clean_text:
                                options.append({
                                    "label": label.upper(),
                                    "text": clean_text
                                })
                        break
                
                if not options:
                    continue
                
                # Extract question text (everything before the first option)
                first_opt = re.search(r'\(?[A-Ea-e]\)?[\.\)]', block, re.IGNORECASE)
                q_text = block[:first_opt.start()].strip() if first_opt else block.strip()
                q_text = re.sub(r'\s+', ' ', q_text).strip()
                
                if q_text and len(options) >= 2:
                    questions.append({
                        "question_text": q_text,
                        "choices": [
                            {
                                "choice_text": o["text"],
                                "is_correct": (o["label"] == correct_answer)
                            }
                            for o in options
                        ]
                    })
            
            if questions:
                break  # Stop if we found questions with this pattern
    
    return questions

@exam_router.delete("/{exam_id}")
async def delete_exam(exam_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    await exam.delete()
    return {"message": "Exam deleted successfully"}

@exam_router.put("/{exam_id}", response_model=ExamResponse)
async def update_exam(exam_id: int, exam_data: ExamCreate):
    db_exam = await Exam.find_one(Exam.int_id == exam_id)
    if not db_exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Update main fields
    db_exam.title = exam_data.title
    db_exam.description = exam_data.description
    db_exam.time_limit = exam_data.time_limit
    db_exam.attempts_allowed = exam_data.attempts_allowed
    db_exam.randomize_questions = exam_data.randomize_questions
    db_exam.status = exam_data.status
    db_exam.course_id = exam_data.course_id
    
    db_exam.questions = []

    question_id_counter = 1
    choice_id_counter = 1
    
    for i, q_data in enumerate(exam_data.questions):
        new_question = ExamQuestion(
            int_id=question_id_counter,
            question_text=q_data.question_text,
            question_type=q_data.question_type,
            points=q_data.points,
            order=q_data.order or i
        )
        question_id_counter += 1

        for choice_data in q_data.choices:
            new_choice = ExamChoice(
                int_id=choice_id_counter,
                choice_text=choice_data.choice_text,
                is_correct=choice_data.is_correct
            )
            choice_id_counter += 1
            new_question.choices.append(new_choice)
            
        db_exam.questions.append(new_question)
    
    await db_exam.save()
    return ExamResponse(**format_exam_response(db_exam))

@exam_router.get("/{exam_id}/attempts", response_model=List[ExamAttemptResponse])
async def get_exam_attempts(exam_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        return []
        
    res = []
    for a in exam.attempts:
        user = await User.find_one(User.int_id == a.student_id)
        
        attempt_res = a.model_dump()
        attempt_res["id"] = a.int_id
        attempt_res["exam_id"] = exam.int_id
        attempt_res["student_name"] = user.name if user else "Unknown"
        attempt_res["student_email"] = user.email if user else "Unknown"
        res.append(attempt_res)
        
    return res

@exam_router.get("/{exam_id}/attempts/{attempt_id}", response_model=ExamAttemptDetailResponse)
async def get_attempt_details(exam_id: int, attempt_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    attempt = next((a for a in exam.attempts if a.int_id == attempt_id), None)
    if not attempt:
        raise HTTPException(status_code=404, detail="Attempt not found")
    
    user = await User.find_one(User.int_id == attempt.student_id)
    
    res = attempt.model_dump()
    res["id"] = attempt.int_id
    res["exam_id"] = exam.int_id
    res["student_name"] = user.name if user else "Unknown"
    res["student_email"] = user.email if user else "Unknown"
    
    formatted_exam = format_exam_response(exam)
    res["exam"] = formatted_exam
    
    formatted_answers = []
    for ans in attempt.answers:
        ans_dict = ans.model_dump()
        ans_dict["id"] = ans.int_id
        ans_dict["question_id"] = ans.question_int_id
        question = next((q for q in formatted_exam.get("questions", []) if q["id"] == ans.question_int_id), None)
        ans_dict["question"] = question
        formatted_answers.append(ans_dict)
        
    res["answers"] = formatted_answers
    
    return res
