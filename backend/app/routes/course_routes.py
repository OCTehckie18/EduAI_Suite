from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from typing import Optional
from app.models.course import Course
from app.models.student import Student
from app.models.assignment import Assignment
from app.models.submission import Submission
from app.models.announcement import Announcement
from app.models.lesson import Lesson
from app.models.resource import Resource
from app.models.exam import Exam
from app.models.slido import PresentationAssignment
from app.schemas.course import CourseCreate, CourseResponse, CourseUpdate
from app.utils.file_uploads import save_optional_upload
from app.utils.auth import get_current_user
from app.models.user import User
import random
import string
import PyPDF2
import docx
import io
import re

course_router = APIRouter(prefix="/courses", tags=["Courses"])


def _same_identity(left: Optional[str], right: Optional[str]) -> bool:
    return bool(left and right and left.strip().casefold() == right.strip().casefold())


def _teacher_identity(user: User) -> str:
    return (user.name or user.email).strip()


async def _course_is_visible(course: Course, current_user: User) -> bool:
    if current_user.role == "admin":
        return True
    if current_user.role == "teacher":
        return _same_identity(course.teacher_name, _teacher_identity(current_user))
    if current_user.role == "student":
        return await Student.find_one(
            Student.email == current_user.email,
            Student.course_id == course.int_id,
        ) is not None
    return False


async def _get_visible_course(course_id: int, current_user: User) -> Course:
    course = await Course.find_one(Course.int_id == course_id)
    if not course or not await _course_is_visible(course, current_user):
        # Do not disclose whether another teacher's classroom exists.
        raise HTTPException(status_code=404, detail="Course not found")
    return course

async def _enrich_course(course: Course) -> dict:
    """Return course as dict with live-computed students count and progress."""
    student_count = await Student.find(Student.course_id == course.int_id).count()

    # Progress = what % of all possible submissions have been made
    # (submissions received / total slots = students × assignments)
    assignments = await Assignment.find(Assignment.course_id == course.int_id).to_list()
    total_slots = student_count * len(assignments)
    if total_slots > 0:
        submission_ids = [a.int_id for a in assignments]
        received = await Submission.find(
            {"assignment_id": {"$in": submission_ids}}
        ).count()
        progress = round(min(received / total_slots * 100, 100), 1)
    else:
        progress = 0.0

    d = course.model_dump()
    d["id"] = course.int_id # Map int_id back to id for the frontend
    d["students"] = student_count
    d["progress"] = progress
    d["color"] = d.get("color") or "#264796"
    d["description"] = d.get("description") or ""
    return d


@course_router.get("/", response_model=list[CourseResponse])
async def get_courses(current_user: User = Depends(get_current_user)):
    courses = await Course.find_all().to_list()
    if current_user.role == "teacher":
        courses = [
            course for course in courses
            if _same_identity(course.teacher_name, _teacher_identity(current_user))
        ]
    elif current_user.role == "student":
        enrolled = await Student.find(Student.email == current_user.email).to_list()
        enrolled_ids = {student.course_id for student in enrolled}
        courses = [course for course in courses if course.int_id in enrolled_ids]
    elif current_user.role != "admin":
        courses = []
    return [await _enrich_course(c) for c in courses]


@course_router.get("/{course_id}", response_model=CourseResponse)
async def get_course(course_id: int, current_user: User = Depends(get_current_user)):
    course = await _get_visible_course(course_id, current_user)
    return await _enrich_course(course)

def generate_unique_code():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))

@course_router.post("/extract_details")
async def extract_course_details(file: UploadFile = File(...)):
    contents = await file.read()
    text = ""
    
    if file.filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(contents))
        raw_text = ""
        for page in reader.pages:
            raw_text += page.extract_text() + "\n"
        # Normalize text: collapse multiple spaces and newlines to avoid PDF-specific layout issues
        text = raw_text  # normalize newlines

    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(contents))
        # Extract text from paragraphs
        paragraphs_text = "\n".join([para.text for para in doc.paragraphs])
        # Extract text from tables (many course plans use tables for metadata)
        tables_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tables_text += "\n" + cell.text
        text = paragraphs_text + "\n" + tables_text
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX.")

    # Parsing logic based on the provided sample image structure
    # Expected fields: Course Code, Course Name, Course Instructor(s) Name, Programme(s)
    
    # regex matches for labels in the document
    details = {
        "code": "",
        "name": "",
        "teacher_name": "",
        "programmes": ""
    }   
    
    # Simple regex based on the image format
    # Use re.IGNORECASE and allow for various separators (colon, spaces, etc.)
    
    # Course Code: SCI202-4L
    # Be greedy with characters typically used in codes
    code_match = re.search(r"Course\s*Code\s*[:\-]?\s*([A-Z0-9\-]+)", text, re.IGNORECASE)
    if code_match: details["code"] = code_match.group(1).strip()
    
    # Course Name: Hypothesis Testing
    name_match = re.search(r"Course Name[:\s]+([^:]+?)(?= Course Code| Course Type| Method| Hours| Credits|$)", text, re.IGNORECASE)
    if not name_match:
        name_match = re.search(r"Course Name[:\s]+([^\n\r]+)", text, re.IGNORECASE)
    if name_match: details["name"] = name_match.group(1).strip()
    

    #  Instructor (handles multiple formats)
    instructor_match = re.search(
        r"(?:Course\s+Instructor\(s\)\s+Name|Instructor)[:\s#]+(.+)",
        text,
        re.IGNORECASE
    )

    if instructor_match:
        details["teacher_name"] = instructor_match.group(1).strip()


    #  Programmes → Description (ROBUST FIX)

    lines = raw_text.split("\n") if file.filename.endswith(".pdf") else paragraphs_text.split("\n")

    desc_lines = []
    capture = False

    for line in lines:
        l = line.strip().lower()

        # START condition
        if "offered programme" in l:
            capture = True
            continue

        # STOP condition
        if any(x in l for x in [
            "course learning outcome",
            "outcome reference",
            "course credit",
            "unit",
            "syllabus"
        ]):
            break

        if capture:
            if line.strip():
                desc_lines.append(line.strip())

    details["description"] = " ".join(desc_lines).strip()

    return details


@course_router.post("/", response_model=CourseResponse, status_code=status.HTTP_201_CREATED)
async def create_course(
    code: str = Form(...),
    name: str = Form(...),
    batch: str = Form(...),
    color: str = Form(...),
    description: str = Form(""),
    enrollment_code: Optional[str] = Form(None),
    teacher_name: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    current_user: User = Depends(get_current_user),
):
    if current_user.role not in ("teacher", "admin"):
        raise HTTPException(status_code=403, detail="Only teachers can create classrooms")
    course_plan_path = save_optional_upload(file, "courses")
    # If no enrollment_code is provided by the system/user, generate a unique 6-char key
    final_code = enrollment_code if enrollment_code else generate_unique_code()
    
    new_course = Course(
        code=code,
        name=name,
        batch=batch,
        students=0,
        progress=0.0,
        color=color,
        description=description,
        enrollment_code=final_code,
        teacher_name=_teacher_identity(current_user) if current_user.role == "teacher" else teacher_name,
        course_plan_path=course_plan_path
    )
    await new_course.assign_id()
    await new_course.insert()
    
    return await _enrich_course(new_course)

@course_router.put("/{course_id}", response_model=CourseResponse)
async def update_course(course_id: int, course_update: CourseUpdate, current_user: User = Depends(get_current_user)):
    course = await _get_visible_course(course_id, current_user)
    
    update_data = course_update.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(course, key, value)
    
    await course.save()
    return await _enrich_course(course)

@course_router.delete("/{course_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_course(course_id: int, current_user: User = Depends(get_current_user)):
    course = await _get_visible_course(course_id, current_user)
    
    # Fetch and delete complex relationships
    exams = await Exam.find(Exam.course_id == course_id).to_list()
    for exam in exams:
        await exam.delete()
        
    pas = await PresentationAssignment.find(PresentationAssignment.course_id == course_id).to_list()
    for pa in pas:
        await pa.delete()

    # Delete all child records
    await Announcement.find(Announcement.course_id == course_id).delete()
    await Assignment.find(Assignment.course_id == course_id).delete()
    await Student.find(Student.course_id == course_id).delete()
    await Lesson.find(Lesson.course_id == course_id).delete()
    await Resource.find(Resource.course_id == course_id).delete()
    
    await course.delete()
    return None
